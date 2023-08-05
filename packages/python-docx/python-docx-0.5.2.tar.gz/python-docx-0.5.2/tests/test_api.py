# encoding: utf-8

"""
Test suite for the docx.api module
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

import pytest

from docx.api import Document
from docx.enum.text import WD_BREAK
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.package import Package
from docx.parts.document import DocumentPart, InlineShapes
from docx.parts.numbering import NumberingPart
from docx.parts.styles import StylesPart
from docx.table import Table
from docx.text import Paragraph, Run

from .unitutil import (
    instance_mock, class_mock, method_mock, property_mock, var_mock
)


class DescribeDocument(object):

    def it_opens_a_docx_on_construction(self, init_fixture):
        docx_, open_ = init_fixture
        document = Document(docx_)
        open_.assert_called_once_with(docx_)
        assert isinstance(document, Document)

    def it_can_open_a_docx_file(self, open_fixture):
        docx_, Package_, package_, document_part_ = open_fixture
        document_part, package = Document._open(docx_)
        Package_.open.assert_called_once_with(docx_)
        assert document_part is document_part
        assert package is package_

    def it_opens_default_template_if_no_file_provided(
            self, Package_, default_docx_):
        Document._open(None)
        Package_.open.assert_called_once_with(default_docx_)

    def it_should_raise_if_not_a_Word_file(self, Package_, package_, docx_):
        package_.main_document.content_type = 'foobar'
        with pytest.raises(ValueError):
            Document._open(docx_)

    def it_can_add_a_heading(self, add_heading_fixture):
        document, add_paragraph_, p_, text, level, style = add_heading_fixture
        p = document.add_heading(text, level)
        add_paragraph_.assert_called_once_with(text, style)
        assert p is p_

    def it_should_raise_on_heading_level_out_of_range(self, document):
        with pytest.raises(ValueError):
            document.add_heading(level=-1)
        with pytest.raises(ValueError):
            document.add_heading(level=10)

    def it_can_add_an_empty_paragraph(self, add_empty_paragraph_fixture):
        document, document_part_, p_ = add_empty_paragraph_fixture
        p = document.add_paragraph()
        document_part_.add_paragraph.assert_called_once_with()
        assert p is p_

    def it_can_add_a_paragraph_of_text(self, add_text_paragraph_fixture):
        document, text, p_, r_ = add_text_paragraph_fixture
        p = document.add_paragraph(text)
        p.add_run.assert_called_once_with()
        r_.add_text.assert_called_once_with(text)

    def it_can_add_a_styled_paragraph(self, add_styled_paragraph_fixture):
        document, style, p_ = add_styled_paragraph_fixture
        p = document.add_paragraph(style=style)
        assert p.style == style

    def it_can_add_a_page_break(self, add_page_break_fixture):
        document, document_part_, p_, r_ = add_page_break_fixture
        p = document.add_page_break()
        document_part_.add_paragraph.assert_called_once_with()
        p_.add_run.assert_called_once_with()
        r_.add_break.assert_called_once_with(WD_BREAK.PAGE)
        assert p is p_

    def it_can_add_a_picture(self, add_picture_fixture):
        (document, image_path, width, height, inline_shapes_, expected_width,
         expected_height, picture_) = add_picture_fixture
        picture = document.add_picture(image_path, width, height)
        inline_shapes_.add_picture.assert_called_once_with(image_path)
        assert picture.width == expected_width
        assert picture.height == expected_height
        assert picture is picture_

    def it_can_add_a_table(self, add_table_fixture):
        document, rows, cols, style, document_part_, expected_style, table_ = (
            add_table_fixture
        )
        table = document.add_table(rows, cols, style)
        document_part_.add_table.assert_called_once_with(rows, cols)
        assert table.style == expected_style
        assert table == table_

    def it_provides_access_to_the_document_inline_shapes(self, document):
        body = document.inline_shapes
        assert body is document._document_part.inline_shapes

    def it_provides_access_to_the_document_paragraphs(
            self, paragraphs_fixture):
        document, paragraphs_ = paragraphs_fixture
        paragraphs = document.paragraphs
        assert paragraphs is paragraphs_

    def it_provides_access_to_the_document_tables(self, tables_fixture):
        document, tables_ = tables_fixture
        tables = document.tables
        assert tables is tables_

    def it_can_save_the_package(self, save_fixture):
        document, package_, file_ = save_fixture
        document.save(file_)
        package_.save.assert_called_once_with(file_)

    def it_provides_access_to_the_numbering_part(self, num_part_get_fixture):
        document, document_part_, numbering_part_ = num_part_get_fixture
        numbering_part = document.numbering_part
        document_part_.part_related_by.assert_called_once_with(RT.NUMBERING)
        assert numbering_part is numbering_part_

    def it_creates_numbering_part_on_first_access_if_not_present(
            self, num_part_create_fixture):
        document, NumberingPart_, document_part_, numbering_part_ = (
            num_part_create_fixture
        )
        numbering_part = document.numbering_part
        NumberingPart_.new.assert_called_once_with()
        document_part_.relate_to.assert_called_once_with(
            numbering_part_, RT.NUMBERING
        )
        assert numbering_part is numbering_part_

    def it_provides_access_to_the_styles_part(self, styles_part_get_fixture):
        document, document_part_, styles_part_ = styles_part_get_fixture
        styles_part = document.styles_part
        document_part_.part_related_by.assert_called_once_with(RT.STYLES)
        assert styles_part is styles_part_

    def it_creates_styles_part_on_first_access_if_not_present(
            self, styles_part_create_fixture):
        document, StylesPart_, document_part_, styles_part_ = (
            styles_part_create_fixture
        )
        styles_part = document.styles_part
        StylesPart_.new.assert_called_once_with()
        document_part_.relate_to.assert_called_once_with(
            styles_part_, RT.STYLES
        )
        assert styles_part is styles_part_

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[0, 1, 2, 5, 9])
    def add_heading_fixture(self, request, document, add_paragraph_, p_):
        level = request.param
        text = 'Spam vs. Bacon'
        style = 'Title' if level == 0 else 'Heading%d' % level
        return document, add_paragraph_, p_, text, level, style

    @pytest.fixture
    def add_empty_paragraph_fixture(self, document, document_part_, p_):
        return document, document_part_, p_

    @pytest.fixture
    def add_page_break_fixture(self, document, document_part_, p_, r_):
        return document, document_part_, p_, r_

    @pytest.fixture
    def add_paragraph_(self, request, p_):
        return method_mock(
            request, Document, 'add_paragraph', return_value=p_
        )

    @pytest.fixture(params=[
        (None, None,  200,  100),
        (1000, 500,  1000,  500),
        (2000, None, 2000, 1000),
        (None, 2000, 4000, 2000),
    ])
    def add_picture_fixture(
            self, request, Document_inline_shapes_, inline_shapes_):
        width, height, expected_width, expected_height = request.param
        document = Document()
        image_path_ = instance_mock(request, str, name='image_path_')
        picture_ = inline_shapes_.add_picture.return_value
        picture_.width, picture_.height = 200, 100
        return (
            document, image_path_, width, height, inline_shapes_,
            expected_width, expected_height, picture_
        )

    @pytest.fixture
    def add_styled_paragraph_fixture(self, document, p_):
        style = 'foobaresque'
        return document, style, p_

    @pytest.fixture(params=[None, 'LightShading-Accent1', 'foobar'])
    def add_table_fixture(self, request, document, document_part_, table_):
        rows, cols = 4, 2
        style = expected_style = request.param
        return (
            document, rows, cols, style, document_part_, expected_style,
            table_
        )

    @pytest.fixture
    def add_text_paragraph_fixture(self, document, p_, r_):
        text = 'foobar\rbarfoo'
        return document, text, p_, r_

    @pytest.fixture
    def default_docx_(self, request):
        return var_mock(request, 'docx.api._default_docx_path')

    @pytest.fixture
    def Document_inline_shapes_(self, request, inline_shapes_):
        return property_mock(
            request, Document, 'inline_shapes', return_value=inline_shapes_
        )

    @pytest.fixture
    def document(self, open_):
        return Document()

    @pytest.fixture
    def document_part_(self, request, p_, paragraphs_, table_, tables_):
        document_part_ = instance_mock(
            request, DocumentPart, content_type=CT.WML_DOCUMENT_MAIN
        )
        document_part_.add_paragraph.return_value = p_
        document_part_.add_table.return_value = table_
        document_part_.paragraphs = paragraphs_
        document_part_.tables = tables_
        return document_part_

    @pytest.fixture
    def docx_(self, request):
        return instance_mock(request, str)

    @pytest.fixture
    def init_fixture(self, docx_, open_):
        return docx_, open_

    @pytest.fixture
    def inline_shapes_(self, request):
        return instance_mock(request, InlineShapes)

    @pytest.fixture
    def num_part_create_fixture(
            self, document, NumberingPart_, document_part_, numbering_part_):
        document_part_.part_related_by.side_effect = KeyError
        return document, NumberingPart_, document_part_, numbering_part_

    @pytest.fixture
    def num_part_get_fixture(self, document, document_part_, numbering_part_):
        document_part_.part_related_by.return_value = numbering_part_
        return document, document_part_, numbering_part_

    @pytest.fixture
    def NumberingPart_(self, request, numbering_part_):
        NumberingPart_ = class_mock(request, 'docx.api.NumberingPart')
        NumberingPart_.new.return_value = numbering_part_
        return NumberingPart_

    @pytest.fixture
    def numbering_part_(self, request):
        return instance_mock(request, NumberingPart)

    @pytest.fixture
    def open_(self, request, document_part_, package_):
        return method_mock(
            request, Document, '_open',
            return_value=(document_part_, package_)
        )

    @pytest.fixture
    def open_fixture(self, docx_, Package_, package_, document_part_):
        return docx_, Package_, package_, document_part_

    @pytest.fixture
    def p_(self, request, r_):
        p_ = instance_mock(request, Paragraph)
        p_.add_run.return_value = r_
        return p_

    @pytest.fixture
    def Package_(self, request, package_):
        Package_ = class_mock(request, 'docx.api.Package')
        Package_.open.return_value = package_
        return Package_

    @pytest.fixture
    def package_(self, request, document_part_):
        package_ = instance_mock(request, Package)
        package_.main_document = document_part_
        return package_

    @pytest.fixture
    def paragraphs_(self, request):
        return instance_mock(request, list)

    @pytest.fixture
    def paragraphs_fixture(self, document, paragraphs_):
        return document, paragraphs_

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, Run)

    @pytest.fixture
    def save_fixture(self, request, open_, package_):
        file_ = instance_mock(request, str)
        document = Document()
        return document, package_, file_

    @pytest.fixture
    def StylesPart_(self, request, styles_part_):
        StylesPart_ = class_mock(request, 'docx.api.StylesPart')
        StylesPart_.new.return_value = styles_part_
        return StylesPart_

    @pytest.fixture
    def styles_part_(self, request):
        return instance_mock(request, StylesPart)

    @pytest.fixture
    def styles_part_create_fixture(
            self, document, StylesPart_, document_part_, styles_part_):
        document_part_.part_related_by.side_effect = KeyError
        return document, StylesPart_, document_part_, styles_part_

    @pytest.fixture
    def styles_part_get_fixture(self, document, document_part_, styles_part_):
        document_part_.part_related_by.return_value = styles_part_
        return document, document_part_, styles_part_

    @pytest.fixture
    def table_(self, request):
        return instance_mock(request, Table, style=None)

    @pytest.fixture
    def tables_(self, request):
        return instance_mock(request, list)

    @pytest.fixture
    def tables_fixture(self, document, tables_):
        return document, tables_
